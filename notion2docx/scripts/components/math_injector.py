import os
import re
import pypandoc
import docx
from docx.oxml.ns import qn

def pypandoc_convert(html_path, docx_path):
    """將 HTML 使用 pandoc 轉成 DOCX"""
    print(f"  [Pandoc] Converting {html_path} to {docx_path}")
    pypandoc.convert_file(html_path, 'docx', outputfile=docx_path, extra_args=['--wrap=none'])

def extract_omml_from_docx(docx_path):
    """
    開啟 pandoc 轉出的 DOCX，解析其中所有的 paragraphs
    並依據 `[ID: MATH_...]` 標籤，把下一段的 OMML 節點取出來，存入字典。
    """
    doc = docx.Document(docx_path)
    math_dict = {}
    current_id = None
    
    for p in doc.paragraphs:
        text = p.text.strip()
        # 如果這是一個 ID 段落
        if text.startswith("[ID: MATH_") and text.endswith("]"):
            current_id = text[5:-1] # 取得 MATH_001
            continue
            
        # 如果前一段是 ID，那這一段或這個 paragraph 裡面的 oMath 就是我們要的
        if current_id:
            # 尋找這個段落內的所有 <m:oMath> (忽略 oMathPara 避免無效的區塊嵌套)
            oMaths = p._p.xpath('.//m:oMath')
            if oMaths:
                import copy
                # 複製該節點 (深拷貝) 保留所有的 XML 屬性與命名空間
                math_dict[current_id] = copy.deepcopy(oMaths[0])
            current_id = None # 重置
            
    print(f"  [Extract] Successfully extracted {len(math_dict)} equations.")
    return math_dict

def inject_omml_into_docx(main_docx_path, math_dict, output_path):
    """
    開啟由 win32com 轉出的帶有原生樣式的 main_docx
    將內文裡的 `[MATH_PLACEHOLDER: MATH_001]` 取代為字典裡對應的 OMML 節點
    """
    doc = docx.Document(main_docx_path)
    
    # 建立一個正則表達式，用來配對 [MATH_PLACEHOLDER: MATH_xxx]
    placeholder_pattern = re.compile(r"\[MATH_PLACEHOLDER:\s*(MATH_\d+)\]")
    
    # 全文掃描：段落
    replace_count = 0
    for p in doc.paragraphs:
        matches = placeholder_pattern.findall(p.text)
        if matches:
            for math_id in matches:
                # 找到對應的 oMath 節點
                if math_id in math_dict:
                    oMath_node = math_dict[math_id]
                    # 尋找該段落下的文字 runs，找出帶有 placeholder 的 run 並取代
                    for run in p.runs:
                        if f"[MATH_PLACEHOLDER: {math_id}]" in run.text:
                            # 清除 placeholder 文字
                            run.text = run.text.replace(f"[MATH_PLACEHOLDER: {math_id}]", "")
                            # 把 oMath 節點作為 run 的兄弟節點插入 (避免放在 w:r 裡面導致 Word 解析錯誤產生方塊)
                            import copy
                            run._r.addnext(copy.deepcopy(oMath_node))
                            replace_count += 1
                                
    # 全文掃描：表格 (Notion 轉換多半是表格)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    matches = placeholder_pattern.findall(p.text)
                    if matches:
                        for math_id in matches:
                            if math_id in math_dict:
                                oMath_node = math_dict[math_id]
                                for run in p.runs:
                                    if f"[MATH_PLACEHOLDER: {math_id}]" in run.text:
                                        run.text = run.text.replace(f"[MATH_PLACEHOLDER: {math_id}]", "")
                                        # 把 oMath 節點作為 run 的兄弟節點插入
                                        import copy
                                        run._r.addnext(copy.deepcopy(oMath_node))
                                        replace_count += 1

    # 存檔
    doc.save(output_path)
    print(f"  [Inject] Injected {replace_count} equations into main document.")
    return replace_count

def process_equations(math_html_path, main_docx_path, final_docx_path):
    """
    統合 pandoc 與 python-docx 進行方程式轉換與替換
    """
    if not os.path.exists(math_html_path):
        print("  [Math Injector] math_library.html not found, skipping equation injection.")
        return False
        
    math_docx_path = math_html_path.replace(".html", ".docx")
    
    try:
        # 1. 將包含方程式的 HTML 交由 pandoc 轉換成 DOCX (以取得 OMML 解析成果)
        pypandoc_convert(math_html_path, math_docx_path)
        
        # 2. 開啟產出的 DOCX，將 OMML 元素提煉出來存入 Dictionary
        math_dict = extract_omml_from_docx(math_docx_path)
        
        # 3. 如果有提煉出任何方程式，就對主文件進行開刀縫合
        if math_dict:
            inject_omml_into_docx(main_docx_path, math_dict, final_docx_path)
            return True
        else:
            # 沒有提取到任何方程式，直接把主文件複製為最終文件
            import shutil
            shutil.copy2(main_docx_path, final_docx_path)
            return True
            
    except Exception as e:
        print(f"Error during math injection process: {e}")
        return False
